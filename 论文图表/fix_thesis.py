"""
论文修改脚本 —— 将论文中的技术描述修正为与实际系统一致。
实际系统: FastAPI + SQLAlchemy + SQLite + Vue3/Vanilla JS + Scrapy
"""
from docx import Document
from docx.shared import Pt, RGBColor
import copy, re, os

SRC = r"C:\Users\kang\OneDrive\桌面\B站ACG视频数据统计分析系统设计与实现(1)(2).docx"
DST = r"C:\Users\kang\OneDrive\桌面\B站ACG视频数据统计分析系统设计与实现_修改版.docx"

doc = Document(SRC)

# ── 辅助函数 ──
def replace_text(obj, old, new):
    """在 paragraph 或 table cell 中替换文本，保留格式"""
    if hasattr(obj, 'runs') and obj.runs:
        full = obj.text
        if old not in full:
            return 0
        # 构建 run 到文本范围的映射
        runs_data = []
        pos = 0
        for r in obj.runs:
            start = pos
            end = pos + len(r.text)
            runs_data.append((r, start, end))
            pos = end
        full_text = ''.join(r.text for r in obj.runs)

        if old not in full_text:
            return 0

        # 简单策略: 如果 old 完全在某一个 run 内，直接替换
        count = 0
        for r, start, end in runs_data:
            if old in r.text:
                r.text = r.text.replace(old, new)
                count += 1
        if count > 0:
            return count

        # 跨 run 策略: 重建所有 run 文本
        idx = full_text.find(old)
        if idx < 0:
            return 0
        new_full = full_text.replace(old, new)
        # 按比例分配回各 run
        if len(full_text) > 0:
            ratio = len(new_full) / len(full_text)
        else:
            ratio = 1.0
        runs_data_sort = sorted([(r, s) for r, s, e in runs_data], key=lambda x: x[1])
        current_pos = 0
        for r, s in runs_data_sort:
            if current_pos >= len(new_full):
                r.text = ""
                continue
            old_len = len(r.text)
            new_len = max(1, int(old_len * ratio))
            end_pos = min(current_pos + new_len, len(new_full))
            r.text = new_full[current_pos:end_pos]
            current_pos = end_pos
        return 1

    return 0


def replace_in_document(doc, old, new):
    """在整个文档中替换文本"""
    count = 0
    for para in doc.paragraphs:
        count += replace_text(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += replace_text(para, old, new)
    return count


def find_paragraph_containing(doc, keyword):
    """查找包含关键词的段落"""
    for i, para in enumerate(doc.paragraphs):
        if keyword in para.text:
            return i, para
    return None, None


# ═══════════════════════════════════════════════════════════
#  第一轮: 核心术语替换 (Django → FastAPI)
# ═══════════════════════════════════════════════════════════
print("=" * 60)
print("  论文修改脚本 — 修正技术描述与实际系统一致")
print("=" * 60)

replacements = [
    # ── 后端框架 ──
    ("Django REST framework", "FastAPI + SQLAlchemy"),
    ("Django框架", "FastAPI框架"),
    ("Django + Django REST framework + Redis", "FastAPI + SQLAlchemy + Uvicorn"),
    ("Django + Django REST framework", "FastAPI + SQLAlchemy"),
    ("基于Django框架", "基于FastAPI框架"),
    ("使用Django", "使用FastAPI"),
    ("后端使用Django", "后端使用FastAPI"),
    ("利用Django", "利用FastAPI"),
    ("在Django中", "在FastAPI中"),
    ("、Django", "、FastAPI"),
    ("Django后端", "FastAPI后端"),
    ("Django 后端", "FastAPI 后端"),
    ("Django框架", "FastAPI框架"),
    ("的Django", "的FastAPI"),

    # ── Django (单独出现，但要小心不替换引用文献中的) ──
    ("采用Django", "采用FastAPI"),
    ("如Django", "如FastAPI"),

    # ── 数据库描述 ──
    ("MySQL + Redis 的组合", "SQLite/MySQL 的组合"),
    ("MySQL作为主数据库", "SQLite作为主数据库（同时兼容MySQL）"),
    ("MySQL作为主", "SQLite作为主（同时兼容MySQL）"),
    ("MySQL数据库", "SQLite数据库（兼容MySQL）"),

    # ── Redis 相关 (实际系统未使用) ──
    ("配合Redis缓存热点数据", "配合数据库索引和查询优化"),
    ("同时使用Redis缓存热点数据", "同时使用数据库索引优化查询"),
    ("Redis缓存", "本地缓存"),

    # ── 数据采集 ──
    ("BeautifulSoup4解析HTML结构", "解析API返回的JSON数据"),
    ("BeautifulSoup解析B站ACG视频页面的HTML结构", "解析B站API返回的JSON数据结构"),
    ("使用BeautifulSoup库解析B站ACG视频页面的HTML结构", "解析B站热门视频API返回的JSON数据"),
    ("通过requests库发送HTTP请求获取排行榜页面HTML", "通过Scrapy发送HTTP请求调用B站热门视频API"),

    # ── 端口和路径 ──
    ("localhost:3001", "localhost:8080"),
    ("端口3001", "端口8080"),
    ("port 3001", "port 8080"),
    ("localhost:5173", "localhost:8080"),

    # ── API 路径 ──
    ("/api/videos/list", "/api/videos"),
    ("/api/videos/search", "/api/videos"),
    ("/api/register", "/api/auth/register"),
    ("/api/login", "/api/auth/login"),
    ("/api/user/profile", "/api/auth/me"),
    ("/api/statistics/", "/api/dashboard/"),

    # ── 前端框架修正 (部分章节误写为React) ──
    ("React 18.2.0", "Vue 3.3.13"),
    ("Ant Design 5.12.8", "Element Plus 2.4.3"),
    ("React", "Vue 3"),

    # ── 后端运行时修正 (误写为Node.js) ──
    ("Node.js 24.14.0", "Python 3.12"),
    ("Express 4.18.2", "FastAPI 0.104.1"),
    ("https://nodejs.org/", "https://www.python.org/"),
    ("https://expressjs.com/", "https://fastapi.tiangolo.com/"),
    ("轻量级Web框架", "高性能异步Web框架"),

    # ── SQLite → 保留 (系统实际使用SQLite) ──
    ("bcryptjs 2.4.3", "passlib[bcrypt] 1.7.4"),
    ("axios 1.6.0", "httpx 0.25.2"),

    # ── DataV (未实际使用) ──
    ("ECharts和DataV实现可视化大屏", "ECharts实现可视化大屏"),
    ("ElementPlus组件库、ECharts可视化库、DataV大屏组件", "ElementPlus组件库、ECharts可视化库"),
    ("ECharts + DataV", "ECharts"),

    # ── 图片代理 → meta referrer ──
    ("图片代理服务", "图片防盗链处理"),
    ("后端代理路由实现 （backend/src/app.js）", "HTML meta标签实现（no-referrer）"),
    ("前端使用代理 （frontend/src/pages/Videos.jsx）", "前端video-card组件中使用cover_url"),

    # ── 数据库 (SQLite) ──
    ("SQLite 3.45.0", "SQLite (内嵌)"),
    ("https://www.sqlite.org/", "https://www.sqlite.org/"),
    ("嵌入式关系型数据库", "关系型数据库"),

    # ── 关键词 ──
    ("数据可视化；Django", "数据可视化；FastAPI"),

    # ── other fixes ──
    ("MVVM 或 MTV", "分层架构"),
    ("DataV实现", "ECharts实现"),
]

for old, new in replacements:
    n = replace_in_document(doc, old, new)
    if n > 0:
        print(f"  [OK] '{old[:50]}' → '{new[:50]}' ({n}处)")

# ═══════════════════════════════════════════════════════════
#  第二轮: 修改/重写特定段落
# ═══════════════════════════════════════════════════════════
print("\n--- 段落级修改 ---")

# 找到并标记所有 "Django" 残留
for i, para in enumerate(doc.paragraphs):
    if "Django" in para.text and "参考文献" not in para.text and "[Django]" not in para.text:
        print(f"  [残留] P{i}: {para.text[:80]}...")

# 找到并标记所有 "React" 残留
for i, para in enumerate(doc.paragraphs):
    if "React" in para.text:
        print(f"  [残留React] P{i}: {para.text[:80]}...")

# ═══════════════════════════════════════════════════════════
#  保存
# ═══════════════════════════════════════════════════════════
doc.save(DST)
print(f"\n{'=' * 60}")
print(f"  修改版已保存至: {DST}")
print(f"  请用 Word 打开查看，以下是仍需手动调整的内容：")
print(f"{'=' * 60}")
print("""
【仍需手动修改的内容】

1. 摘要(中文+英文)：
   - "后端基于Django框架" → "后端基于FastAPI框架"
   - 关键词 "Django" → "FastAPI"
   - 英文 "Django framework" → "FastAPI framework"

2. 第2.6节 Django章节：
   - 整节替换为 FastAPI 介绍
   - FastAPI：基于 Starlette + ASGI 的高性能 Python Web 框架
   - 强调：自动生成 OpenAPI 文档、Pydantic 类型校验、依赖注入

3. 第4.4节 数据库设计：
   - E-R 图应包含 5 张表: users, videos, comments, user_favorites, video_stats_history
   - 补充 comments 表(评论,支持嵌套回复)
   - 补充 user_favorites 表(用户收藏)
   - 补充 video_stats_history 表(历史统计快照)
   - videos 表补充字段: share_count, reply_count, duration, description, tags, bilibili_url, author_id

4. 第5章 系统实现：
   - 5.2数据采集: 删除 BeautifulSoup HTML解析描述，改为 API JSON 解析
   - 5.3图片代理: 删除代理服务描述，改为 <meta referrer="no-referrer"> 方案
   - 5.4视频播放: 删除嵌入式播放器描述，改为跳转B站原视频

5. 第6章 部署说明：
   - 后端: Python + uvicorn 启动 (非 Node.js/Express)
   - 前端: Vite 开发服务器 或 FastAPI 静态文件服务
   - 端口: 后端 8080，前端 5173

6. 补充缺失功能描述(建议在第3章/第5章补充)：
   - 评论系统(支持嵌套回复)
   - 用户收藏功能
   - 个人中心(头像/昵称/密码修改)
   - 数据看板(互动率分析、时长分析)
   - 可视化大屏(全屏模式、实时时钟)
""")
