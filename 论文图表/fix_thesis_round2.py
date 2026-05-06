"""
第二轮精细修改：处理残留 Django 引用，重写特定段落
"""
from docx import Document
from docx.shared import Pt, RGBColor
import re, os

DST = r"C:\Users\kang\OneDrive\桌面\B站ACG视频数据统计分析系统设计与实现_修改版.docx"

doc = Document(DST)

# ── 辅助函数 ──
def set_para_text(para, new_text):
    """完全替换段落文本，保留第一个run的格式"""
    if para.runs:
        fmt = para.runs[0].font
        for r in para.runs:
            r.text = ""
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

def replace_in_para(para, old, new):
    """在段落中替换文本"""
    full = para.text
    if old not in full:
        return False
    # 合并所有run然后重建
    full_text = ''.join(r.text for r in para.runs)
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    for r in para.runs:
        r.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)
    return True


# ═══════════════════════════════════════════════════════════
# 1. 英文摘要修复
# ═══════════════════════════════════════════════════════════
for para in doc.paragraphs:
    if "Django" in para.text:
        if "Hadoop" in para.text or "参考文献" in para.text:
            # 引用的文献，不改
            continue
        replace_in_para(para, "Django", "FastAPI")
        print(f"  [FIX] 残留Django: {para.text[:80]}...")

# ═══════════════════════════════════════════════════════════
# 2. 重写第2.6节 — Django → FastAPI
# ═══════════════════════════════════════════════════════════
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "2.6 Django" or para.text.strip() == "2.6 FastAPI":
        para.runs[0].text = "2.6 FastAPI"
        print(f"  [FIX] 章节标题: 2.6 FastAPI")

    if "Django 是一款基于 Python 编程语言的开源、高级、全栈式 Web 应用开发框架" in para.text:
        # 完全重写这个段落为 FastAPI 介绍
        new_text = (
            "FastAPI 是一款基于 Python 编程语言的现代、高性能 Web 应用开发框架，"
            "由 Sebastián Ramírez 于 2018 年创建，基于 Starlette 异步框架和 Pydantic 数据校验库构建。"
            "FastAPI 以高性能（媲美 Node.js 和 Go）、极简代码、自动生成交互式 API 文档、"
            "原生异步支持为核心优势，被广泛应用于数据 API 服务、微服务架构、机器学习模型部署等场景。"
        )
        set_para_text(para, new_text)
        print(f"  [REWRITE] P{i}: FastAPI 介绍段落")

    if "Django 框架采用松耦合、高内聚的模块化设计" in para.text:
        new_text = (
            "FastAPI 框架采用松耦合、高内聚的模块化设计，将 Web 应用的数据模型、路由控制、业务逻辑进行明确分离，"
            "使系统结构更加清晰，便于开发、维护与扩展。其核心架构分为三个层次："
            "Schema 层负责请求/响应数据结构定义与校验（基于 Pydantic）；"
            "Router 层负责路由注册与请求分发；"
            "Dependency 层通过依赖注入机制实现数据库会话管理、认证鉴权等横切关注点。"
            "FastAPI 能够自动生成符合 OpenAPI 3.0 规范的交互式 API 文档（Swagger UI 和 ReDoc），"
            "开发者无需额外编写文档代码，即可获得完整的 API 测试界面。"
            "这种设计有效提升了代码复用率，降低了模块间的依赖关系，特别适合前后端分离架构下的 RESTful API 开发。"
        )
        set_para_text(para, new_text)
        print(f"  [REWRITE] P{i}: FastAPI 架构描述")

# ═══════════════════════════════════════════════════════════
# 3. 修复第2章小结中的技术栈描述
# ═══════════════════════════════════════════════════════════
for i, para in enumerate(doc.paragraphs):
    if "Scrapy 爬虫框架、MySQL 数据库、ECharts 可视化库、Requests 请求库以及" in para.text:
        replace_in_para(para, "以及 Django Web 框架", "以及 FastAPI Web 框架")
        print(f"  [FIX] P{i}: 小结技术栈")

    if "Django 框架则承担后端服务、路由控制、业务逻辑与前后端交互的核心任务" in para.text:
        replace_in_para(para,
            "Django 框架则承担后端服务、路由控制、业务逻辑与前后端交互的核心任务",
            "FastAPI 框架则承担后端服务、路由控制、业务逻辑与前后端交互的核心任务")
        print(f"  [FIX] P{i}: 小结 Django→FastAPI")

    # 修正 Python 章节中对 Django 的引用
    if "在Web开发领域有Django、Flask、" in para.text:
        replace_in_para(para, "Django、", "")
        print(f"  [FIX] P{i}: 移除 Django 列举")

# ═══════════════════════════════════════════════════════════
# 4. 修复数据采集层描述中的 "Django管理命令"
# ═══════════════════════════════════════════════════════════
for i, para in enumerate(doc.paragraphs):
    if "通过Django自定义命令调用" in para.text:
        replace_in_para(para, "Django自定义命令", "Python脚本命令")
        print(f"  [FIX] P{i}: Django自定义命令 → Python脚本命令")
    if "通过Django管理命令或Celery定时任务触发" in para.text:
        replace_in_para(para, "通过Django管理命令或Celery定时任务触发", "通过Python脚本或系统定时任务（cron/scheduler）触发")
        print(f"  [FIX] P{i}: 爬虫触发方式")

# ═══════════════════════════════════════════════════════════
# 5. 修复第7章总结中的技术路线描述
# ═══════════════════════════════════════════════════════════
for i, para in enumerate(doc.paragraphs):
    if "后端使用Django + Django REST framework + Redis" in para.text:
        replace_in_para(para,
            "后端使用Django + Django REST framework + Redis",
            "后端使用FastAPI + SQLAlchemy + Uvicorn")
        print(f"  [FIX] P{i}: 总结技术路线")

# ═══════════════════════════════════════════════════════════
# 6. 修复表 4.1 模块接口表中的 API 路径
# ═══════════════════════════════════════════════════════════
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "Django" in para.text:
                    replace_in_para(para, "Django", "FastAPI")

# ═══════════════════════════════════════════════════════════
# 7. 修正视频表字段描述
# ═══════════════════════════════════════════════════════════
# (论文中视频表缺少很多字段，尽可能修正)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text.strip() == "view_count":
                    replace_in_para(para, "view_count", "play_count")

# ═══════════════════════════════════════════════════════════
# 8. 补充章节 2.3: 添加 Scrapy 框架专门章节说明
# ═══════════════════════════════════════════════════════════
# (论文目前没有 Scrapy 独立章节，只在 2.1 Python 中一带而过)
# 查找 2.2 节的位置，考虑在适合位置补入
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("2.3") and "MySQL" in para.text:
        print(f"  [INFO] 2.3 MySQL 在 P{i}, 建议在此之前插入 2.2 Scrapy 章节(需手动)")

# ═══════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════
doc.save(DST)
print(f"\n  [OK] 第二轮修改完成，保存至: {DST}")
print(f"\n  ⚠ 以下内容仍需手动在 Word 中修改：")
print(f"  1. 第4.4节: 更新数据库 E-R 图和表结构（3表→5表）")
print(f"  2. 第5章: 更新代码示例（Python/FastAPI 替代 Node.js/Express）")
print(f"  3. 表5.1/5.2: 更新开发环境表（移除 Node.js/Express/React/Ant Design）")
print(f"  4. 第3.2节: 补充评论、收藏、个人中心等功能需求")
print(f"  5. 添加缺失的 Scrapy 独立章节 (2.2)")
